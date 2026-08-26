"""Download public contract templates and build local RAG knowledge.

Source: 国家市场监督管理总局合同示范文本库
https://htsfwb.samr.gov.cn/List
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import subprocess
import time
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any
from urllib.parse import quote

import httpx
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE_BASE = "https://htsfwb.samr.gov.cn"
LIST_API = f"{SOURCE_BASE}/api/content/SearchTemplates"
DOWNLOAD_API = f"{SOURCE_BASE}/api/File/DownTemplate"
CATALOG_FILENAME = "catalog.json"
HEADERS = {
    "User-Agent": "Mozilla/5.0 LegalAgentWorkbench/0.1",
    "Referer": f"{SOURCE_BASE}/List",
}


RISK_KEYWORDS: dict[str, tuple[str, str, tuple[str, ...]]] = {
    "payment_acceptance": ("medium", "付款与验收节点", ("付款", "支付", "价款", "验收", "发票")),
    "unlimited_liability": ("high", "赔偿责任边界", ("赔偿", "全部损失", "违约责任", "损失赔偿", "责任")),
    "data_security": ("high", "数据与信息安全", ("个人信息", "数据", "隐私", "泄露", "安全")),
    "ip_ownership": ("high", "知识产权归属", ("知识产权", "著作权", "专利", "成果", "商标")),
    "auto_renewal": ("medium", "期限与续约", ("续约", "自动续期", "服务期限", "合同期限")),
    "jurisdiction": ("medium", "争议解决与管辖", ("争议", "管辖", "法院", "仲裁")),
    "confidentiality": ("medium", "保密义务", ("保密", "商业秘密", "秘密信息")),
    "termination_notice": ("medium", "解除与终止通知", ("解除", "终止", "通知", "提前")),
    "force_majeure": ("low", "不可抗力", ("不可抗力", "不能预见", "不能避免")),
}


@dataclass
class CorpusItem:
    id: str
    title: str
    department: str
    region: str
    year: str
    is_local: bool
    type: int
    detail_url: str
    download_url: str
    raw_path: str
    markdown_path: str
    chars: int
    clauses: int
    status: str
    sha256: str = ""
    duplicate_of: str = ""
    error: str = ""


class SourceClient:
    """HTTPX source client with a process-wide curl fallback for CDN TLS quirks."""

    def __init__(self) -> None:
        self.httpx_client = httpx.Client(headers=HEADERS, timeout=30.0, follow_redirects=True)
        self.transport = "httpx"
        self.fallback_reason = ""

    def __enter__(self) -> "SourceClient":
        return self

    def __exit__(self, *_args: object) -> None:
        self.httpx_client.close()

    def get_json(self, url: str, *, params: dict[str, Any]) -> dict[str, Any]:
        if self.transport == "httpx":
            try:
                return get_json_with_retry(self.httpx_client, url, params=params)
            except RuntimeError as exc:
                self._switch_to_curl(exc)
        content = curl_request(url, params=params, timeout=60)
        try:
            payload = json.loads(content)
        except json.JSONDecodeError as exc:
            raise RuntimeError(f"curl catalog response is not valid JSON: {exc}") from exc
        if not isinstance(payload, dict):
            raise RuntimeError("curl catalog response is not a JSON object")
        return payload

    def get_file(self, url: str, *, params: dict[str, Any]) -> bytes:
        if self.transport == "httpx":
            last_error: Exception | None = None
            for attempt in range(3):
                try:
                    response = self.httpx_client.get(url, params=params)
                    response.raise_for_status()
                    return response.content
                except (httpx.TransportError, httpx.TimeoutException, httpx.HTTPStatusError) as exc:
                    last_error = exc
                    if attempt < 2:
                        time.sleep(1.0 + attempt)
            self._switch_to_curl(last_error or RuntimeError("httpx download failed"))
        return curl_request(url, params=params, timeout=180)

    def _switch_to_curl(self, error: Exception) -> None:
        if not shutil.which("curl"):
            raise RuntimeError(f"httpx failed and curl is unavailable: {error}") from error
        self.transport = "curl"
        self.fallback_reason = str(error)
        print(f"source transport switched to curl: {self.fallback_reason}", flush=True)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--limit", type=int, default=100, help="目标下载份数（示范文本库存量支持 500+）")
    parser.add_argument("--resume", action="store_true", help="已存在的 docx 不重复下载，只补缺")
    parser.add_argument("--offline", action="store_true", help="不访问网络，仅用 catalog/manifest 重建本地语料")
    parser.add_argument("--output", default="data/common_contracts")
    parser.add_argument("--knowledge-output", default=".lawbench/knowledge/common_contract_corpus.json")
    args = parser.parse_args()

    if args.limit < 1:
        parser.error("--limit must be at least 1")

    output = ROOT / args.output
    raw_dir = output / "raw_docx"
    md_dir = output / "markdown"
    raw_dir.mkdir(parents=True, exist_ok=True)
    md_dir.mkdir(parents=True, exist_ok=True)
    (output / "README.md").write_text(_readme(), encoding="utf-8")

    previous_items = load_previous_items(output)
    catalog_source = "remote"
    source_error = ""
    with SourceClient() as client:
        if args.offline:
            templates = load_cached_templates(output)
            catalog_source = "cache"
        else:
            try:
                templates = fetch_templates(max(args.limit * 2, args.limit + 60), client=client)
                (output / CATALOG_FILENAME).write_text(
                    json.dumps({"generated_at": time.time(), "templates": templates}, ensure_ascii=False, indent=2) + "\n",
                    encoding="utf-8",
                )
            except Exception as exc:
                if not args.resume:
                    raise
                templates = load_cached_templates(output)
                catalog_source = "cache"
                source_error = str(exc)
                print(f"remote catalog unavailable; using {len(templates)} cached templates: {source_error}", flush=True)
        if not templates:
            raise RuntimeError("no contract templates available from remote catalog or local cache")

        items: list[CorpusItem] = []
        knowledge: list[dict[str, Any]] = []
        seen_content: dict[str, str] = {}
        success_count = 0
        for idx, template in enumerate(templates, start=1):
            item, entries = process_template(
                client,
                template,
                raw_dir=raw_dir,
                md_dir=md_dir,
                index=idx,
                resume=args.resume or args.offline,
                previous_item=previous_items.get(str(template["Id"])),
                allow_download=not args.offline,
            )
            if item.status == "ok":
                duplicate_of = seen_content.get(item.sha256)
                if duplicate_of:
                    item.status = "duplicate"
                    item.duplicate_of = duplicate_of
                    item.error = f"exact content duplicate of {duplicate_of}"
                else:
                    seen_content[item.sha256] = item.id
                    knowledge.extend(entries)
            items.append(item)
            if item.status == "ok":
                success_count += 1
                if success_count % 25 == 0 or success_count == args.limit:
                    print(f"processed {success_count}/{args.limit} contracts", flush=True)
            if success_count >= args.limit:
                break
            if not args.offline:
                time.sleep(0.05)

    manifest = {
        "source": "国家市场监督管理总局合同示范文本库",
        "source_url": f"{SOURCE_BASE}/List",
        "generated_at": time.time(),
        "requested_limit": args.limit,
        "target_met": success_count >= args.limit,
        "catalog_source": catalog_source,
        "catalog_size": len(templates),
        "network_transport": client.transport,
        "transport_fallback_reason": client.fallback_reason,
        "source_error": source_error,
        "downloaded": sum(1 for item in items if item.status == "ok"),
        "duplicates": sum(1 for item in items if item.status == "duplicate"),
        "failed": sum(1 for item in items if item.status == "failed"),
        "knowledge_entries": len(knowledge),
        "items": [asdict(item) for item in items],
    }
    (output / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    (output / "CORPUS_REPORT.md").write_text(render_report(manifest), encoding="utf-8")
    knowledge_path = ROOT / args.knowledge_output
    knowledge_path.parent.mkdir(parents=True, exist_ok=True)
    knowledge_path.write_text(json.dumps(knowledge, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(
        json.dumps(
            {
                k: manifest[k]
                for k in ("requested_limit", "target_met", "downloaded", "duplicates", "failed", "knowledge_entries")
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0 if manifest["target_met"] else 2


def fetch_templates(limit: int, *, client: SourceClient | None = None) -> list[dict[str, Any]]:
    templates: list[dict[str, Any]] = []
    seen_ids: set[str] = set()
    owned_client = client is None
    source = client or SourceClient()
    try:
        for is_local in (False, True):
            page = 1
            while len(templates) < limit:
                payload = source.get_json(LIST_API, params={"loc": str(is_local).lower(), "p": page})
                data = payload.get("Data") or []
                if not data:
                    break
                for template in data:
                    template_id = str(template.get("Id") or "")
                    if not template_id or template_id in seen_ids:
                        continue
                    seen_ids.add(template_id)
                    templates.append(template)
                if page >= int(payload.get("TotalPage") or 1):
                    break
                page += 1
                time.sleep(0.05)
                if len(templates) >= limit:
                    break
    finally:
        if owned_client:
            source.httpx_client.close()
    return templates[:limit]


def curl_request(url: str, *, params: dict[str, Any], timeout: int) -> bytes:
    executable = shutil.which("curl")
    if not executable:
        raise RuntimeError("curl is unavailable")
    command = [
        executable,
        "--silent",
        "--show-error",
        "--fail",
        "--location",
        "--retry",
        "3",
        "--retry-all-errors",
        "--retry-delay",
        "1",
        "--connect-timeout",
        "20",
        "--max-time",
        str(timeout),
        "--header",
        f"User-Agent: {HEADERS['User-Agent']}",
        "--header",
        f"Referer: {HEADERS['Referer']}",
        "--get",
    ]
    for key, value in params.items():
        command.extend(["--data-urlencode", f"{key}={value}"])
    command.append(url)
    try:
        result = subprocess.run(command, check=False, capture_output=True, timeout=timeout + 30)
    except (OSError, subprocess.SubprocessError) as exc:
        raise RuntimeError(f"curl request failed: {exc}") from exc
    if result.returncode != 0:
        stderr = result.stderr.decode("utf-8", errors="replace").strip()
        raise RuntimeError(f"curl request failed with exit {result.returncode}: {stderr[:300]}")
    return result.stdout


def get_json_with_retry(
    client: SourceClient,
    url: str,
    *,
    params: dict[str, Any],
    attempts: int = 3,
) -> dict[str, Any]:
    last_error: Exception | None = None
    for attempt in range(attempts):
        try:
            response = client.get(url, params=params)
            response.raise_for_status()
            payload = response.json()
            if not isinstance(payload, dict):
                raise ValueError("catalog response is not a JSON object")
            return payload
        except (httpx.TransportError, httpx.TimeoutException, httpx.HTTPStatusError, json.JSONDecodeError, ValueError) as exc:
            last_error = exc
            if attempt + 1 < attempts:
                time.sleep(1.0 + attempt)
    raise RuntimeError(f"catalog request failed after {attempts} attempts: {last_error}") from last_error


def load_previous_items(output: Path) -> dict[str, dict[str, Any]]:
    path = output / "manifest.json"
    if not path.exists():
        return {}
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}
    items = payload.get("items") if isinstance(payload, dict) else []
    if not isinstance(items, list):
        return {}
    return {str(item.get("id")): item for item in items if isinstance(item, dict) and item.get("id")}


def load_cached_templates(output: Path) -> list[dict[str, Any]]:
    catalog_path = output / CATALOG_FILENAME
    if catalog_path.exists():
        try:
            payload = json.loads(catalog_path.read_text(encoding="utf-8"))
            templates = payload.get("templates") if isinstance(payload, dict) else None
            if isinstance(templates, list) and templates:
                return [item for item in templates if isinstance(item, dict) and item.get("Id")]
        except (OSError, json.JSONDecodeError):
            pass
    return [cached_item_to_template(item) for item in load_previous_items(output).values()]


def cached_item_to_template(item: dict[str, Any]) -> dict[str, Any]:
    return {
        "Id": str(item.get("id") or ""),
        "Title": str(item.get("title") or ""),
        "Department": str(item.get("department") or ""),
        "Region": str(item.get("region") or ""),
        "PublishedOn": str(item.get("year") or ""),
        "IsLocal": bool(item.get("is_local")),
        "Type": int(item.get("type") or 0),
    }


def process_template(
    client: httpx.Client,
    template: dict[str, Any],
    *,
    raw_dir: Path,
    md_dir: Path,
    index: int,
    resume: bool = False,
    previous_item: dict[str, Any] | None = None,
    allow_download: bool = True,
) -> tuple[CorpusItem, list[dict[str, Any]]]:
    template_id = str(template["Id"])
    title = str(template.get("Title") or template_id)
    safe = f"{template_id}_{safe_name(title)}"
    raw_path = resume_path(previous_item, "raw_path", raw_dir, raw_dir / f"{safe}.docx")
    md_path = resume_path(previous_item, "markdown_path", md_dir, md_dir / f"{safe}.md")
    detail_url = f"{SOURCE_BASE}/View?id={template_id}"
    download_url = f"{DOWNLOAD_API}?id={quote(template_id)}&type=1"
    try:
        has_reusable_file = resume and is_word_file(raw_path)
        if not has_reusable_file:
            if not allow_download:
                raise FileNotFoundError(f"cached Word file unavailable: {raw_path}")
            content = client.get_file(DOWNLOAD_API, params={"id": template_id, "type": 1})
            if not is_word_payload(content):
                raise ValueError(f"downloaded payload is not a Word file: bytes={len(content)}")
            raw_path.write_bytes(content)
        text = extract_docx_text(raw_path)
        md_path.write_text(markdown_for(template, text), encoding="utf-8")
        entries = build_knowledge_entries(template, text, detail_url)
        item = CorpusItem(
            id=template_id,
            title=title,
            department=str(template.get("Department") or ""),
            region=str(template.get("Region") or ""),
            year=str(template.get("PublishedOn") or ""),
            is_local=bool(template.get("IsLocal")),
            type=int(template.get("Type") or 0),
            detail_url=detail_url,
            download_url=download_url,
            raw_path=str(raw_path.relative_to(ROOT)),
            markdown_path=str(md_path.relative_to(ROOT)),
            chars=len(text),
            clauses=len(entries),
            status="ok",
            sha256=hashlib.sha256(raw_path.read_bytes()).hexdigest(),
        )
        return item, entries
    except Exception as exc:
        item = CorpusItem(
            id=template_id,
            title=title,
            department=str(template.get("Department") or ""),
            region=str(template.get("Region") or ""),
            year=str(template.get("PublishedOn") or ""),
            is_local=bool(template.get("IsLocal")),
            type=int(template.get("Type") or 0),
            detail_url=detail_url,
            download_url=download_url,
            raw_path=str(raw_path.relative_to(ROOT)),
            markdown_path=str(md_path.relative_to(ROOT)),
            chars=0,
            clauses=0,
            status="failed",
            error=str(exc),
        )
        return item, []


def resume_path(
    previous_item: dict[str, Any] | None,
    key: str,
    expected_dir: Path,
    fallback: Path,
) -> Path:
    if not previous_item or not previous_item.get(key):
        return fallback
    candidate = (ROOT / str(previous_item[key])).resolve()
    try:
        candidate.relative_to(expected_dir.resolve())
    except ValueError:
        return fallback
    return candidate


def is_word_payload(content: bytes) -> bool:
    return len(content) >= 512 and (content.startswith(b"PK\x03\x04") or content.startswith(b"\xd0\xcf\x11\xe0"))


def is_word_file(path: Path) -> bool:
    if not path.exists() or path.stat().st_size < 512:
        return False
    try:
        return is_word_payload(path.read_bytes()[:512])
    except OSError:
        return False


def extract_docx_text(path: Path) -> str:
    errors: list[str] = []
    try:
        doc = Document(str(path))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = clean_extracted_text("\n".join(parts))
        if text:
            return text
        errors.append("python-docx returned empty text")
    except Exception as exc:  # noqa: BLE001 - legacy .doc needs a platform converter
        errors.append(f"python-docx: {exc}")

    converters = [
        ("textutil", ["-convert", "txt", "-stdout", "--", str(path)]),
        ("antiword", [str(path)]),
    ]
    for executable, arguments in converters:
        command = shutil.which(executable)
        if not command:
            continue
        try:
            result = subprocess.run(
                [command, *arguments],
                check=False,
                capture_output=True,
                timeout=120,
            )
        except (OSError, subprocess.SubprocessError) as exc:
            errors.append(f"{executable}: {exc}")
            continue
        text = clean_extracted_text(result.stdout.decode("utf-8", errors="replace"))
        if result.returncode == 0 and text:
            return text
        stderr = result.stderr.decode("utf-8", errors="replace").strip()
        errors.append(f"{executable}: exit={result.returncode} {stderr[:160]}")
    raise RuntimeError("unable to extract Word text; " + "; ".join(errors))


def clean_extracted_text(text: str) -> str:
    normalized = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    return "\n".join(line.rstrip() for line in normalized.splitlines()).strip()


def build_knowledge_entries(template: dict[str, Any], text: str, detail_url: str) -> list[dict[str, Any]]:
    contract_type = infer_contract_type(str(template.get("Title") or ""))
    chunks = split_clauses(text)
    entries: list[dict[str, Any]] = []
    for idx, chunk in enumerate(chunks, start=1):
        risk_type, risk_level, clause_type = classify_clause(chunk)
        entries.append(
            {
                "id": f"samr_{template['Id']}_{idx:03d}",
                "title": f"{template.get('Title')} · {clause_title(chunk)}",
                "body": chunk[:1200],
                "contract_type": contract_type,
                "clause_type": clause_type,
                "risk_type": risk_type,
                "risk_level": risk_level,
                "source": f"samr_contract_template:{template['Id']}",
                "tags": [
                    str(template.get("Department") or ""),
                    str(template.get("Region") or ""),
                    str(template.get("PublishedOn") or ""),
                    detail_url,
                ],
            }
        )
    return entries


def split_clauses(text: str) -> list[str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    chunks: list[list[str]] = []
    current: list[str] = []
    heading = re.compile(r"^(第[一二三四五六七八九十百]+条|\\d+[\\.、]|[一二三四五六七八九十]+、)")
    for line in lines:
        if heading.match(line) and current:
            chunks.append(current)
            current = [line]
        else:
            current.append(line)
        if sum(len(x) for x in current) > 900:
            chunks.append(current)
            current = []
    if current:
        chunks.append(current)
    return ["\n".join(chunk) for chunk in chunks if len("".join(chunk)) >= 30][:24]


def classify_clause(text: str) -> tuple[str, str, str]:
    for risk_type, (level, clause_type, keywords) in RISK_KEYWORDS.items():
        if any(keyword in text for keyword in keywords):
            return risk_type, level, clause_type
    return "general", "low", "general"


def infer_contract_type(title: str) -> str:
    mapping = [
        ("保密", "NDA"),
        ("采购", "procurement"),
        ("买卖", "sales"),
        ("租赁", "lease"),
        ("服务", "service"),
        ("能源", "service"),
        ("物业", "service"),
        ("建设", "construction"),
        ("旅游", "consumer"),
        ("养老", "consumer"),
        ("劳动", "employment"),
        ("培训", "consumer"),
        ("预付", "consumer"),
    ]
    for keyword, value in mapping:
        if keyword in title:
            return value
    return "general"


def clause_title(text: str) -> str:
    first = text.splitlines()[0].strip()
    return first[:60] or "合同条款"


def safe_name(value: str) -> str:
    return re.sub(r"[^\w.\-]+", "_", value, flags=re.UNICODE).strip("._")[:80] or "contract"


def markdown_for(template: dict[str, Any], text: str) -> str:
    return "\n".join(
        [
            f"# {template.get('Title')}",
            "",
            "- 来源：国家市场监督管理总局合同示范文本库",
            f"- ID：`{template.get('Id')}`",
            f"- 部门：{template.get('Department') or ''}",
            f"- 地区：{template.get('Region') or ''}",
            f"- 年份：{template.get('PublishedOn') or ''}",
            f"- 链接：{SOURCE_BASE}/View?id={template.get('Id')}",
            "",
            "## 正文",
            "",
            text,
            "",
        ]
    )


def render_report(manifest: dict[str, Any]) -> str:
    return "\n".join(
        [
            "# Common Contract Corpus Report",
            "",
            f"- Source: {manifest['source']}",
            f"- Source URL: {manifest['source_url']}",
            f"- Downloaded: {manifest['downloaded']}",
            f"- Exact-content duplicates skipped: {manifest.get('duplicates', 0)}",
            f"- Failed: {manifest['failed']}",
            f"- Knowledge entries: {manifest['knowledge_entries']}",
            f"- Target met: {manifest.get('target_met', False)}",
            f"- Catalog source: {manifest.get('catalog_source', 'remote')}",
            "",
            "This corpus is built from public contract model texts for local RAG and benchmark development.",
            "",
        ]
    )


def _readme() -> str:
    return """# Common Contract Corpus

This folder stores public contract model texts downloaded from 国家市场监督管理总局合同示范文本库.

- `raw_docx/`: original Word files
- `markdown/`: extracted text for inspection
- `catalog.json`: cached source catalog for reliable resume/offline rebuild
- `manifest.json`: source metadata
- `CORPUS_REPORT.md`: corpus summary

The source sometimes serves legacy `.doc` binaries with a `.docx` filename. The builder first uses
`python-docx`, then falls back to macOS `textutil` or `antiword` when available.
Exact-content duplicates are retained in the manifest for provenance but excluded from the effective
contract count and generated RAG knowledge.

Generated by `scripts/build_common_contract_corpus.py`.
"""


if __name__ == "__main__":
    raise SystemExit(main())
