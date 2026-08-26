"""File-backed contract upload and ingestion."""

from __future__ import annotations

import base64
import json
import os
import re
import tempfile
import time
from pathlib import Path
from typing import Any
from uuid import uuid4

from legalworkbench.paths import uploads_dir
from legalworkbench.secure_storage import (
    secure_read_text,
    secure_write_bytes,
    secure_write_text,
)

SUPPORTED_TEXT_EXTENSIONS = {".txt", ".md", ".markdown"}
SUPPORTED_BINARY_EXTENSIONS = {".pdf", ".docx"}


class ContractDocumentStore:
    """Manage uploaded contract files and extracted text."""

    def __init__(self, cwd: str | Path | None = None) -> None:
        self.cwd = Path(cwd or Path.cwd()).resolve()
        self.root = uploads_dir(self.cwd)
        self.index_path = self.root / "documents.json"

    def save_text(
        self,
        *,
        filename: str,
        text: str,
        source: str = "web",
        tenant_id: str = "local",
        user_id: str = "",
    ) -> dict[str, Any]:
        document_id = f"doc_{uuid4().hex[:10]}"
        safe_name = _safe_filename(filename)
        path = self.root / f"{document_id}.md.enc"
        secure_write_text(
            path,
            text.rstrip() + "\n",
            cwd=self.cwd,
            purpose="uploaded-contract-text",
        )
        record = {
            "document_id": document_id,
            "tenant_id": tenant_id,
            "user_id": user_id,
            "filename": safe_name,
            "path": str(path),
            "source": source,
            "chars": len(text),
            "created_at": time.time(),
            "status": "ready",
        }
        self._append(record)
        return record

    def save_base64(
        self,
        *,
        filename: str,
        content_base64: str,
        source: str = "web_upload",
        tenant_id: str = "local",
        user_id: str = "",
    ) -> dict[str, Any]:
        raw = base64.b64decode(content_base64)
        return self.save_bytes(
            filename=filename,
            data=raw,
            source=source,
            tenant_id=tenant_id,
            user_id=user_id,
        )

    def save_bytes(
        self,
        *,
        filename: str,
        data: bytes,
        source: str = "web_upload",
        tenant_id: str = "local",
        user_id: str = "",
    ) -> dict[str, Any]:
        suffix = Path(filename).suffix.lower()
        raw = data
        raw_path = self.root / f"raw_{uuid4().hex[:10]}{suffix or '.bin'}"
        secure_write_bytes(
            raw_path,
            raw,
            cwd=self.cwd,
            purpose="uploaded-contract-original",
        )
        if suffix in SUPPORTED_TEXT_EXTENSIONS:
            text = raw.decode("utf-8", errors="replace")
        elif suffix in SUPPORTED_BINARY_EXTENSIONS:
            text = _extract_document_bytes(raw, suffix=suffix, root=self.root)
        else:
            text = f"# Unsupported document preview\n\n文件 `{filename}` 已接收，但当前解析器支持 txt/md/pdf/docx。"
        record = self.save_text(
            filename=filename,
            text=text,
            source=source,
            tenant_id=tenant_id,
            user_id=user_id,
        )
        record["raw_path"] = str(raw_path)
        self._replace(record)
        return record

    def list(
        self, limit: int = 50, *, tenant_id: str | None = None
    ) -> list[dict[str, Any]]:
        if not self.index_path.exists():
            return []
        try:
            rows = json.loads(secure_read_text(self.index_path, cwd=self.cwd))
        except (json.JSONDecodeError, UnicodeDecodeError):
            return []
        if not isinstance(rows, list):
            return []
        if tenant_id is not None:
            rows = [
                row
                for row in rows
                if str(row.get("tenant_id") or "local") == tenant_id
            ]
        rows.sort(key=lambda item: item.get("created_at", 0), reverse=True)
        return rows[:limit]

    def get(
        self, document_id: str, *, tenant_id: str | None = None
    ) -> dict[str, Any] | None:
        for record in self.list(limit=500, tenant_id=tenant_id):
            if record.get("document_id") == document_id:
                return record
        return None

    def read_text(self, document_id: str, *, tenant_id: str | None = None) -> str:
        record = self.get(document_id, tenant_id=tenant_id)
        if record is None:
            raise FileNotFoundError(document_id)
        return secure_read_text(record["path"], cwd=self.cwd)

    def _append(self, record: dict[str, Any]) -> None:
        rows = self.list(limit=500)
        rows.insert(0, record)
        self._write_index(rows)

    def _replace(self, record: dict[str, Any]) -> None:
        rows = self.list(limit=500)
        rows = [
            record if item.get("document_id") == record.get("document_id") else item
            for item in rows
        ]
        self._write_index(rows)

    def _write_index(self, rows: list[dict[str, Any]]) -> None:
        secure_write_text(
            self.index_path,
            json.dumps(rows, ensure_ascii=False, indent=2) + "\n",
            cwd=self.cwd,
            purpose="uploaded-contract-index",
        )


def _safe_filename(filename: str) -> str:
    name = Path(filename or "contract.md").name
    name = re.sub(r"[^A-Za-z0-9_.\-\u4e00-\u9fff]+", "_", name).strip("._")
    return name or "contract.md"


def extract_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_document_bytes(data: bytes, *, suffix: str, root: Path) -> str:
    """Parse a binary attachment through a private, short-lived plaintext file."""

    fd, temp_name = tempfile.mkstemp(prefix=".parse_", suffix=suffix, dir=str(root))
    temp_path = Path(temp_name)
    try:
        with os.fdopen(fd, "wb") as handle:
            handle.write(data)
            handle.flush()
            os.fsync(handle.fileno())
        return extract_document_text(temp_path)
    finally:
        try:
            temp_path.unlink()
        except OSError:
            pass


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:  # pragma: no cover - optional dependency
        return (
            f"# PDF parse unavailable\n\n缺少 pypdf 依赖，无法解析 `{path.name}`：{exc}"
        )
    reader = PdfReader(str(path))
    parts = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            parts.append(f"## Page {index}\n{text.strip()}")
    if parts:
        return "\n\n".join(parts)
    # 文本层为空 -> 大概率是扫描件（图片型 PDF），走 OCR 扩展点
    ocr_text = _try_ocr_pdf(path)
    if ocr_text:
        return ocr_text
    return (
        f"# Scanned PDF (needs OCR)\n\n未能从 `{path.name}` 抽取文本层，疑似扫描件。"
        "安装 OCR 依赖后重新上传可解析：`pip install rapidocr-onnxruntime pdf2image`。"
    )


def _try_ocr_pdf(path: Path) -> str:
    """扫描件 OCR 扩展点：依赖存在则逐页 OCR，缺依赖返回空由上层标注 needs_ocr。

    刻意选择本地 OCR（rapidocr）而非云端 OCR API：合同扫描件含 PII 与商业秘密，
    与远端 LLM 脱敏同一原则——明文不出本地信任边界。
    """

    try:  # pragma: no cover - optional heavy dependency
        from pdf2image import convert_from_path  # type: ignore
        from rapidocr_onnxruntime import RapidOCR  # type: ignore
    except Exception:
        return ""
    try:  # pragma: no cover - exercised only when OCR deps installed
        engine = RapidOCR()
        parts = []
        for index, image in enumerate(convert_from_path(str(path), dpi=200), start=1):
            import numpy as np  # type: ignore

            result, _ = engine(np.array(image))
            lines = [item[1] for item in (result or [])]
            if lines:
                parts.append(f"## Page {index} (OCR)\n" + "\n".join(lines))
        return "\n\n".join(parts)
    except Exception:
        return ""


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # pragma: no cover - optional dependency
        return f"# DOCX parse unavailable\n\n缺少 python-docx 依赖，无法解析 `{path.name}`：{exc}"
    document = Document(str(path))
    paragraphs = [
        item.text.strip() for item in document.paragraphs if item.text.strip()
    ]
    tables = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                tables.append(" | ".join(cells))
    text = "\n\n".join([*paragraphs, *tables])
    return text or f"# Empty DOCX\n\n未能从 `{path.name}` 抽取文本。"
