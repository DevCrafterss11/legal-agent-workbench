"""Regression tests for the public contract corpus builder."""

from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document

from scripts import build_common_contract_corpus as corpus


def test_safe_name_preserves_chinese_and_removes_path_separators() -> None:
    value = corpus.safe_name("数据委托处理/服务合同（2025 版）")
    assert value == "数据委托处理_服务合同_2025_版"
    assert "/" not in value


def test_extract_docx_text_reads_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "contract.docx"
    document = Document()
    document.add_paragraph("第一条 付款")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "付款节点"
    table.cell(0, 1).text = "验收后"
    document.save(path)

    text = corpus.extract_docx_text(path)

    assert "第一条 付款" in text
    assert "付款节点 | 验收后" in text


def test_extract_docx_text_falls_back_for_legacy_doc(monkeypatch, tmp_path: Path) -> None:
    path = tmp_path / "legacy.docx"
    path.write_bytes(b"\xd0\xcf\x11\xe0" + b"\x00" * 508)

    def broken_document(_path: str):
        raise ValueError("legacy binary Word file")

    completed = subprocess.CompletedProcess(
        args=["textutil"],
        returncode=0,
        stdout="第一条  旧版 Word 合同\r\n正文".encode(),
        stderr=b"",
    )
    monkeypatch.setattr(corpus, "Document", broken_document)
    monkeypatch.setattr(corpus.shutil, "which", lambda name: "/usr/bin/textutil" if name == "textutil" else None)
    monkeypatch.setattr(corpus.subprocess, "run", lambda *args, **kwargs: completed)

    assert corpus.extract_docx_text(path) == "第一条  旧版 Word 合同\n正文"


def test_word_payload_validation_rejects_html() -> None:
    assert corpus.is_word_payload(b"PK\x03\x04" + b"\x00" * 508)
    assert corpus.is_word_payload(b"\xd0\xcf\x11\xe0" + b"\x00" * 508)
    assert not corpus.is_word_payload(b"<html>upstream error</html>" * 100)


def test_resume_path_only_accepts_files_inside_expected_directory(tmp_path: Path, monkeypatch) -> None:
    root = tmp_path / "repo"
    raw_dir = root / "data" / "raw"
    raw_dir.mkdir(parents=True)
    fallback = raw_dir / "stable.docx"
    monkeypatch.setattr(corpus, "ROOT", root)

    accepted = corpus.resume_path({"raw_path": "data/raw/old.docx"}, "raw_path", raw_dir, fallback)
    rejected = corpus.resume_path({"raw_path": "../outside.docx"}, "raw_path", raw_dir, fallback)

    assert accepted == raw_dir / "old.docx"
    assert rejected == fallback
