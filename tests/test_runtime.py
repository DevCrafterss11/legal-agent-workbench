from pathlib import Path
import base64
import json

from typer.testing import CliRunner

from legalworkbench.cli import app
from legalworkbench.feishu_api import FeishuApiError, FeishuDownloadedFile
from legalworkbench.feishu_events import FeishuEventBridge, extract_document_id, normalize_feishu_event_payload
from legalworkbench.feishu_stream import FeishuLongConnectionListener
from legalworkbench.lark_mcp import configure_lark_mcp, lark_mcp_status
from legalworkbench.mcp import _run_async_blocking
from legalworkbench.documents.store import ContractDocumentStore
from legalworkbench.evals import BaselineEvaluator, HumanBenchmarkRunner, format_baseline_table
from legalworkbench.models import ContractClause, KnowledgeEntry, LegalMemory, LegalSkill, RetrievedEvidence, ReviewRun, RiskFinding
from legalworkbench.paths import knowledge_dir, skills_path
from legalworkbench.rag import clear_rag_service_cache, get_rag_service
import legalworkbench.rag.service as rag_service
from legalworkbench.runtime import LegalAgentRuntime
from legalworkbench.secrets import connector_secret
from legalworkbench.store import write_model_list
from legalworkbench.tasks import ReviewTaskQueue, ReviewTaskWorker
from legalworkbench.tools.base import ToolContext
from legalworkbench.tools.rewrite import ClauseRewriterTool
from legalworkbench.web import render_app_html


def test_runtime_review_eval_and_cli(tmp_path: Path) -> None:
    runtime = LegalAgentRuntime(tmp_path)
    paths = runtime.init_samples()
    run = runtime.review(paths["contract"])
    assert run.status in {"completed", "blocked"}
    assert run.findings
    assert run.report_path
    assert run.mcp_context["agent_architecture"]["pattern"] == "supervisor_worker"
    assert "evidence_agent" in run.mcp_context["agent_architecture"]["workers"]
    assert "Agent 执行架构" in run.report_markdown
    assert any(trace.metadata.get("agent") == "evidence_agent" for trace in run.tool_calls)

    result = runtime.benchmark()
    assert result.cases >= 1
    assert result.tool_success_rate == 1.0

    runner = CliRunner()
    cli_result = runner.invoke(app, ["runs", "--cwd", str(tmp_path)])
    assert cli_result.exit_code == 0
    assert run.review_run_id in cli_result.output

    help_result = runner.invoke(app, ["serve", "--help"])
    assert help_result.exit_code == 0
    assert "bind port" in help_result.output.lower()


def test_human_annotated_benchmark_runner(tmp_path: Path) -> None:
    LegalAgentRuntime(tmp_path).init_samples()
    bench_dir = tmp_path / "data" / "human_benchmark"
    contracts_dir = bench_dir / "contracts"
    contracts_dir.mkdir(parents=True)
    contract_file = contracts_dir / "sample.md"
    contract_file.write_text(
        "# 测试 SaaS 协议\n\n"
        "## 1. 赔偿责任\n"
        "乙方应赔偿甲方全部损失，且不设赔偿责任上限。\n",
        encoding="utf-8",
    )
    (bench_dir / "annotations.json").write_text(
        json.dumps(
            {
                "contracts": [
                    {
                        "contract_id": "sample",
                        "title": "测试 SaaS 协议",
                        "contract_type": "SaaS",
                        "scenario": "单条款测试",
                        "file": "contracts/sample.md",
                        "annotations": [
                            {
                                "risk_id": "HR0001",
                                "clause_id": "C001",
                                "clause_title": "赔偿责任",
                                "risk_type": "unlimited_liability",
                                "risk_level": "high",
                                "rationale": "责任无上限。",
                                "expected_suggestion": "设置责任上限。",
                                "evidence_source": "company_policy:liability_v1",
                                "requires_human_review": True,
                            }
                        ],
                    }
                ]
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    result = HumanBenchmarkRunner(tmp_path).run()

    assert result.contracts == 1
    assert result.annotated_risks == 1
    assert result.risk_recall_at_10 == 1.0
    assert result.rule_recall == 1.0


def test_baseline_evaluator_and_cli(tmp_path: Path) -> None:
    LegalAgentRuntime(tmp_path).init_samples()

    rows = BaselineEvaluator(tmp_path).run(dataset="synthetic")

    assert [row.method for row in rows] == ["rule_only", "rag_only", "full_system"]
    assert all(row.dataset == "synthetic" for row in rows)
    assert all(row.expected_risks > 0 for row in rows)
    rule_only = next(row for row in rows if row.method == "rule_only")
    full_system = next(row for row in rows if row.method == "full_system")
    assert full_system.risk_recall_at_10 >= rule_only.risk_recall_at_10
    assert "full_system" in format_baseline_table(rows)

    runner = CliRunner()
    cli_result = runner.invoke(app, ["eval-baseline", "--cwd", str(tmp_path), "--dataset", "synthetic"])
    assert cli_result.exit_code == 0
    assert "Baseline comparison" in cli_result.output
    assert "rule_only" in cli_result.output


def test_lark_mcp_configuration_keeps_secret_local(tmp_path: Path) -> None:
    status = configure_lark_mcp(
        tmp_path,
        app_id="cli_test_app",
        app_secret="test_secret_value",
        tools=["docx.v1.document.rawContent", "task.v2.task.create"],
    )
    assert status["configured"] is True
    assert status["app_id_configured"] is True
    assert status["app_secret_configured"] is True
    assert "test_secret_value" not in str(status)

    secret = connector_secret("feishu_legal_workspace", tmp_path)
    assert secret["APP_SECRET"] == "test_secret_value"

    status_again = lark_mcp_status(tmp_path)
    assert status_again["tools"] == ["docx.v1.document.rawContent", "task.v2.task.create"]
    assert "test_secret_value" not in str(status_again)

    runner = CliRunner()
    cli_result = runner.invoke(app, ["lark-mcp", "--cwd", str(tmp_path)])
    assert cli_result.exit_code == 0
    assert "test_secret_value" not in cli_result.output
    assert "app_secret_configured" in cli_result.output


def test_web_exposes_milvus_and_bge_controls() -> None:
    html = render_app_html()
    assert "RAG / Milvus" in html
    assert "BGE / sentence-transformers" in html
    assert "embeddingProvider" in html


def test_feishu_event_bridge_challenge_and_text_review(tmp_path: Path) -> None:
    LegalAgentRuntime(tmp_path).init_samples()
    bridge = FeishuEventBridge(tmp_path)

    assert bridge.handle({"challenge": "abc"}) == {"challenge": "abc"}
    assert extract_document_id("https://example.feishu.cn/docx/AbCdEf123456") == "AbCdEf123456"
    unauthenticated = bridge.handle({"schema": "2.0", "event": {}})
    assert unauthenticated["ok"] is False
    assert "authentication is not configured" in unauthenticated["error"]

    result = bridge.handle(
        {
            "schema": "2.0",
            "header": {"event_type": "im.message.receive_v1"},
            "event": {
                "sender": {"sender_id": {"open_id": "ou_test"}},
                "message": {
                    "message_id": "msg_test",
                    "chat_id": "",
                    "message_type": "text",
                    "content": '{"text":"## 赔偿责任\\n乙方承担全部损失且不设责任上限。"}',
                },
            },
        },
        trusted_source=True,
    )
    assert result["ok"] is True
    assert result["status"] == "reviewed"
    assert result["review_run_id"].startswith("law_")
    assert "reply" in result


def test_feishu_event_bridge_ignores_duplicate_message(tmp_path: Path) -> None:
    LegalAgentRuntime(tmp_path).init_samples()
    bridge = FeishuEventBridge(tmp_path)
    payload = {
        "schema": "2.0",
        "header": {"event_type": "im.message.receive_v1"},
        "event": {
            "sender": {"sender_id": {"open_id": "ou_test"}},
            "message": {
                "message_id": "msg_duplicate",
                "chat_id": "",
                "message_type": "text",
                "content": '{"text":"## 赔偿责任\\n乙方承担全部损失且不设责任上限。"}',
            },
        },
    }

    first = bridge.handle(payload, trusted_source=True)
    second = bridge.handle(payload, trusted_source=True)

    assert first["status"] == "reviewed"
    assert second["status"] == "ignored"
    assert second["reason"] == "duplicate message"
    assert len(LegalAgentRuntime(tmp_path).store.list_runs()) == 1


def test_feishu_event_bridge_ignores_self_app_message(tmp_path: Path) -> None:
    configure_lark_mcp(tmp_path, app_id="cli_self_app", app_secret="test_secret_value", tools=["im.v1.message.create"])
    bridge = FeishuEventBridge(tmp_path)
    result = bridge.handle(
        {
            "schema": "2.0",
            "header": {"event_type": "im.message.receive_v1"},
            "event": {
                "sender": {"sender_type": "app", "sender_id": {"app_id": "cli_self_app"}},
                "message": {
                    "message_id": "msg_self",
                    "chat_id": "oc_test",
                    "message_type": "text",
                    "content": '{"text":"合同审查完成：law_old\\n\\n完整报告：/tmp/old.md"}',
                },
            },
        },
        trusted_source=True,
    )

    assert result["status"] == "ignored"
    assert result["reason"] == "self app message"
    assert LegalAgentRuntime(tmp_path).store.list_runs() == []


def test_feishu_event_bridge_downloads_and_reviews_file_attachment(tmp_path: Path) -> None:
    from docx import Document
    from io import BytesIO

    LegalAgentRuntime(tmp_path).init_samples()
    bridge = FeishuEventBridge(tmp_path)
    doc = Document()
    doc.add_heading("赔偿责任", level=2)
    doc.add_paragraph("乙方承担全部损失且不设赔偿责任上限。")
    buf = BytesIO()
    doc.save(buf)

    class FakeFeishuApi:
        def download_message_file(self, **kwargs):
            assert kwargs["message_id"] == "msg_file"
            assert kwargs["file_key"] == "file_v3_test"
            return FeishuDownloadedFile(
                filename=kwargs["filename"],
                content=buf.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    bridge.feishu_api = FakeFeishuApi()
    result = bridge.handle(
        {
            "schema": "2.0",
            "header": {"event_type": "im.message.receive_v1"},
            "event": {
                "sender": {"sender_id": {"open_id": "ou_test"}},
                "message": {
                    "message_id": "msg_file",
                    "chat_id": "",
                    "message_type": "file",
                        "content": '{"file_key":"file_v3_test","file_name":"contract.docx"}',
                },
            },
        },
        trusted_source=True,
    )

    assert result["ok"] is True
    assert result["status"] == "reviewed"
    assert result["file_name"] == "contract.docx"
    assert result["review_run_id"].startswith("law_")
    run = LegalAgentRuntime(tmp_path).store.load_run(result["review_run_id"])
    assert run is not None
    assert any(finding.risk_type == "unlimited_liability" for finding in run.findings)
    docs = ContractDocumentStore(tmp_path).list()
    assert docs[0]["source"] == "feishu_file"


def test_feishu_event_bridge_reports_file_download_failure(tmp_path: Path) -> None:
    LegalAgentRuntime(tmp_path).init_samples()
    bridge = FeishuEventBridge(tmp_path)

    class FailingFeishuApi:
        def download_message_file(self, **kwargs):
            raise FeishuApiError("missing permission im:message:file")

    bridge.feishu_api = FailingFeishuApi()
    result = bridge.handle(
        {
            "schema": "2.0",
            "header": {"event_type": "im.message.receive_v1"},
            "event": {
                "sender": {"sender_id": {"open_id": "ou_test"}},
                "message": {
                    "message_id": "msg_file_failure",
                    "chat_id": "",
                    "message_type": "file",
                    "content": '{"file_key":"file_v3_test","file_name":"contract.pdf"}',
                },
            },
        },
        trusted_source=True,
    )

    assert result["ok"] is False
    assert result["status"] == "failed"
    assert result["reason"] == "file download failed"
    assert LegalAgentRuntime(tmp_path).store.list_runs() == []


def test_feishu_reply_is_actionable_for_high_risk_clause(tmp_path: Path) -> None:
    run = ReviewRun(
        review_run_id="law_test_reply",
        contract_path=str(tmp_path / "contract.md"),
        contract_type="general",
        created_at=1.0,
        updated_at=1.0,
        clauses=[
            ContractClause(
                clause_id="C001",
                title="赔偿责任",
                text="乙方承担全部损失且不设赔偿责任上限。",
            )
        ],
        findings=[
            RiskFinding(
                finding_id="F001",
                clause_id="C001",
                clause_title="赔偿责任",
                risk_type="unlimited_liability",
                risk_level="high",
                summary="条款可能要求承担无限责任或过宽赔偿范围。",
                evidence=[
                    RetrievedEvidence(
                        entry_id="risk_unlimited_liability",
                        title="无限责任赔偿风险",
                        source="company_policy:liability_v1",
                        score=9.0,
                        reason="hybrid",
                        body_preview="出现无限责任、全部损失、间接损失、预期利润损失等表述时，应建议设置赔偿责任上限。",
                        risk_type="unlimited_liability",
                        risk_level="high",
                    )
                ],
                suggestion="建议改为：乙方仅对因其违约造成的直接损失承担赔偿责任，累计赔偿总额不超过事故发生前十二个月甲方已实际支付的服务费用；间接损失、预期利润损失、商誉损失、业务中断损失及惩罚性赔偿不纳入赔偿范围。",
                requires_human_review=True,
            )
        ],
    )

    reply = FeishuEventBridge(tmp_path)._build_reply(run)

    assert "原文：乙方承担全部损失且不设赔偿责任上限。" in reply
    assert "风险影响：责任范围没有边界" in reply
    assert "company_policy:liability_v1" in reply
    assert "建议替换为：乙方仅对因其违约造成的直接损失承担赔偿责任" in reply
    assert "累计赔偿总额不超过事故发生前十二个月" in reply


def test_clause_rewriter_prefers_specific_template_over_short_rule(tmp_path: Path) -> None:
    result = ClauseRewriterTool().execute(
        {
            "risk_type": "unlimited_liability",
            "rule_suggestion": "建议设置赔偿责任上限。",
            "memories": [
                LegalMemory(
                    memory_id="mem_short",
                    type="episodic",
                    risk_type="unlimited_liability",
                    summary="旧短建议",
                    approved_advice="建议设置赔偿责任上限。",
                    confidence=0.95,
                )
            ],
            "evidence": [],
        },
        ToolContext(cwd=tmp_path, review_run_id="law_test"),
    )

    assert "累计赔偿总额不超过事故发生前十二个月" in result.output
    assert "间接损失" in result.output


def test_uploaded_docx_can_be_reviewed(tmp_path: Path) -> None:
    from docx import Document

    docx_path = tmp_path / "contract.docx"
    document = Document()
    document.add_heading("赔偿责任", level=2)
    document.add_paragraph("乙方承担全部损失且不设赔偿责任上限。")
    document.save(docx_path)

    encoded = base64.b64encode(docx_path.read_bytes()).decode("ascii")
    store = ContractDocumentStore(tmp_path)
    record = store.save_base64(filename="contract.docx", content_base64=encoded)

    assert record["status"] == "ready"
    assert "赔偿责任" in store.read_text(record["document_id"])
    raw_file = next((tmp_path / ".lawbench" / "uploads").glob("raw_*.docx"))
    assert raw_file.stat().st_mode & 0o777 == 0o600
    assert (tmp_path / ".lawbench").stat().st_mode & 0o777 == 0o700
    assert (tmp_path / ".lawbench" / "uploads").stat().st_mode & 0o777 == 0o700

    run = LegalAgentRuntime(tmp_path).review(record["path"])

    assert run.findings
    assert run.findings[0].risk_type == "unlimited_liability"
    assert "累计赔偿总额不超过事故发生前十二个月" in run.findings[0].suggestion


def test_encrypted_uploaded_contract_can_be_reviewed(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("LEGAL_WORKBENCH_ENCRYPTION_PROVIDER", "env")
    monkeypatch.setenv(
        "LEGAL_WORKBENCH_ENCRYPTION_KEY",
        base64.urlsafe_b64encode(b"e" * 32).decode("ascii"),
    )
    runtime = LegalAgentRuntime(tmp_path)
    runtime.init_samples()
    store = ContractDocumentStore(tmp_path)

    record = store.save_text(
        filename="张三合同.md",
        text="姓名：张三\n## 赔偿责任\n乙方承担全部损失且不设责任上限。",
    )

    persisted = Path(record["path"]).read_bytes()
    assert persisted.startswith(b"LAWBENCH-ENC-v1\n")
    assert "张三".encode() not in persisted
    assert store.read_text(record["document_id"]).startswith("姓名：张三")
    assert runtime.review(record["path"]).status == "completed"


def test_rag_service_factory_reuses_warm_instance(tmp_path: Path, monkeypatch) -> None:
    LegalAgentRuntime(tmp_path).init_samples()
    clear_rag_service_cache()
    calls = {"count": 0}
    real_service = rag_service.LegalRagService

    class CountedLegalRagService(real_service):
        def __init__(self, *args, **kwargs):
            calls["count"] += 1
            super().__init__(*args, **kwargs)

    monkeypatch.setattr(rag_service, "LegalRagService", CountedLegalRagService)

    first = get_rag_service(tmp_path)
    second = get_rag_service(tmp_path)

    assert first is second
    assert calls["count"] == 1


def test_review_reuses_rag_service_across_clause_retrievals(tmp_path: Path, monkeypatch) -> None:
    LegalAgentRuntime(tmp_path).init_samples()
    clear_rag_service_cache()
    calls = {"count": 0}
    real_service = rag_service.LegalRagService

    class CountedLegalRagService(real_service):
        def __init__(self, *args, **kwargs):
            calls["count"] += 1
            super().__init__(*args, **kwargs)

    monkeypatch.setattr(rag_service, "LegalRagService", CountedLegalRagService)
    contract = tmp_path / "multi_clause.md"
    contract.write_text(
        "\n\n".join(
            [
                "## 自动续约\n服务期满后自动续约一年。",
                "## 赔偿责任\n乙方承担全部损失且不设赔偿责任上限。",
                "## 数据安全\n乙方处理客户数据但未约定泄露通知时限。",
            ]
        ),
        encoding="utf-8",
    )

    run = LegalAgentRuntime(tmp_path).review(contract)

    assert run.findings
    assert len(run.clauses) == 3
    assert calls["count"] == 1


def test_rag_service_cache_invalidates_when_knowledge_changes(tmp_path: Path) -> None:
    paths = LegalAgentRuntime(tmp_path).init_samples()
    clear_rag_service_cache()
    first = get_rag_service(tmp_path)
    knowledge_path = paths["knowledge"]
    knowledge_path.write_text(knowledge_path.read_text(encoding="utf-8") + "\n", encoding="utf-8")
    second = get_rag_service(tmp_path)

    assert first is not second


def test_feishu_long_connection_status_and_payload_normalization(tmp_path: Path) -> None:
    configure_lark_mcp(
        tmp_path,
        app_id="cli_test_app",
        app_secret="test_secret_value",
        tools=["im.v1.message.create"],
    )
    status = FeishuLongConnectionListener(tmp_path).status()
    assert status["configured"] is True
    assert status["requires_public_domain"] is False
    assert status["event"] == "im.message.receive_v1"

    payload = normalize_feishu_event_payload(
        {
            "sender": {"sender_id": {"open_id": "ou_test"}},
            "message": {"message_type": "text", "content": '{"text":"合同"}'},
        }
    )
    assert payload["schema"] == "2.0"
    assert payload["header"]["event_type"] == "im.message.receive_v1"


def test_mcp_async_runner_works_inside_existing_event_loop() -> None:
    import asyncio

    async def value() -> str:
        await asyncio.sleep(0)
        return "ok"

    async def outer() -> str:
        return _run_async_blocking(value(), timeout=1)

    assert asyncio.run(outer()) == "ok"


def test_skill_focus_can_promote_rag_evidence_into_finding(tmp_path: Path) -> None:
    runtime = LegalAgentRuntime(tmp_path)
    runtime.init_samples()
    write_model_list(
        knowledge_dir(tmp_path) / "custom.json",
        [
            KnowledgeEntry(
                id="risk_termination_notice",
                title="单方解约通知期风险",
                body="合同允许任一方随时单方解除且没有提前通知期时，法务应要求补充 30 天书面通知和过渡期。",
                contract_type="support",
                clause_type="termination",
                risk_type="termination_notice",
                risk_level="medium",
                source="playbook:support_termination",
                tags=["解约", "通知期"],
            )
        ],
    )
    write_model_list(
        skills_path(tmp_path),
        [
            LegalSkill(
                name="support_contract_review",
                contract_type="support",
                description="审查技术支持合同中的解约和服务边界。",
                focus_clause_types=["termination"],
                risk_rules=["termination_notice"],
                retrieval_top_k=12,
                report_style="risk-first",
                review_playbook=["核对是否存在无通知期单方解约。"],
            )
        ],
    )
    contract = tmp_path / "support.md"
    contract.write_text("## 单方解除\n任一方可以随时单方解除本合同，无需提前通知。", encoding="utf-8")

    run = LegalAgentRuntime(tmp_path).review(contract)

    assert "support_contract_review" in run.selected_skills
    assert run.mcp_context["skill_profile"]["retrieval_top_k"] == 12
    finding = next(item for item in run.findings if item.risk_type == "termination_notice")
    assert "skill_focus" in finding.rule_hits
    assert "Skill 审查策略" in run.report_markdown


def test_task_queue_reviews_uploaded_contract_and_links_report(tmp_path: Path) -> None:
    LegalAgentRuntime(tmp_path).init_samples()
    record = ContractDocumentStore(tmp_path).save_text(
        filename="task-contract.md",
        text="## 赔偿责任\n乙方承担全部损失且不设赔偿责任上限。",
    )
    task = ReviewTaskQueue(tmp_path).add(
        title="队列审查赔偿责任",
        source="test",
        contract_path=record["path"],
        document_id=record["document_id"],
    )
    before = ReviewTaskQueue(tmp_path).summary()

    assert before["remaining"] == 1
    assert before["pending"] == 1
    assert before["next_tasks"][0]["task_id"] == task["task_id"]

    result = ReviewTaskWorker(tmp_path).run_once()
    after = ReviewTaskQueue(tmp_path).summary()

    assert result is not None
    assert result["task_id"] == task["task_id"]
    assert result["status"] == "completed"
    assert result["document_id"] == record["document_id"]
    assert result["review_run_id"].startswith("law_")
    assert Path(result["report_path"]).exists()
    assert after["remaining"] == 0
    assert after["completed"] == 1
    assert after["recent_completed"][0]["review_run_id"] == result["review_run_id"]


def test_task_queue_can_delete_invalid_legacy_failures(tmp_path: Path) -> None:
    queue = ReviewTaskQueue(tmp_path)
    invalid = queue.add(title="旧任务", source="manual")
    valid = queue.add(title="有效失败任务", source="manual", contract_path="/tmp/missing.md")
    queue.update(invalid["task_id"], status="failed", error="contract_path required")
    queue.update(valid["task_id"], status="failed", error="missing file")

    deleted = queue.delete_failed_without_contract()

    assert deleted == 1
    remaining = queue.list()
    assert [item["task_id"] for item in remaining] == [valid["task_id"]]
    assert queue.delete(valid["task_id"]) is True
    assert queue.delete("missing") is False
    assert queue.list() == []
